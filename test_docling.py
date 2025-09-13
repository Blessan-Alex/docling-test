#!/usr/bin/env python3
"""
Comprehensive test script for docling package installation and functionality
Tests various document formats supported by docling
"""

import os
import sys
import tempfile
import subprocess
from pathlib import Path
import requests
from docling.document_converter import DocumentConverter
from docling.datamodel.base_models import InputFormat

def test_pip_install():
    """Test pip installation of docling"""
    print("=" * 60)
    print("TESTING PIP INSTALLATION OF DOCLING")
    print("=" * 60)
    
    try:
        # Test basic pip install
        result = subprocess.run([sys.executable, "-m", "pip", "install", "docling"], 
                              capture_output=True, text=True, check=True)
        print("✅ Basic pip install: SUCCESS")
        print(f"Output: {result.stdout}")
        
        # Test pip install with verbose output
        result = subprocess.run([sys.executable, "-m", "pip", "install", "-v", "docling"], 
                              capture_output=True, text=True, check=True)
        print("✅ Verbose pip install: SUCCESS")
        
        # Test pip install with no-deps (to test core package)
        result = subprocess.run([sys.executable, "-m", "pip", "install", "--no-deps", "docling"], 
                              capture_output=True, text=True, check=True)
        print("✅ No-deps pip install: SUCCESS")
        
        # Test pip install with force-reinstall
        result = subprocess.run([sys.executable, "-m", "pip", "install", "--force-reinstall", "docling"], 
                              capture_output=True, text=True, check=True)
        print("✅ Force-reinstall pip install: SUCCESS")
        
        return True
        
    except subprocess.CalledProcessError as e:
        print(f"❌ Pip install failed: {e}")
        print(f"Error output: {e.stderr}")
        return False

def test_import():
    """Test importing docling modules"""
    print("\n" + "=" * 60)
    print("TESTING DOCLING IMPORTS")
    print("=" * 60)
    
    try:
        import docling
        print("✅ Main docling module: SUCCESS")
        
        from docling.document_converter import DocumentConverter
        print("✅ DocumentConverter import: SUCCESS")
        
        from docling.datamodel.base_models import InputFormat
        print("✅ InputFormat import: SUCCESS")
        
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        print("✅ PdfPipelineOptions import: SUCCESS")
        
        return True
        
    except ImportError as e:
        print(f"❌ Import failed: {e}")
        return False

def create_sample_files():
    """Create sample files in various formats for testing"""
    print("\n" + "=" * 60)
    print("CREATING SAMPLE TEST FILES")
    print("=" * 60)
    
    sample_files = {}
    
    # Create temporary directory
    temp_dir = Path(tempfile.mkdtemp())
    print(f"📁 Created temp directory: {temp_dir}")
    
    # 1. Plain text file
    txt_file = temp_dir / "sample.txt"
    txt_content = """Sample Text Document
====================

This is a sample text document for testing docling.

Features:
- Plain text content
- Multiple paragraphs
- Simple formatting

Testing docling's ability to process basic text files.
"""
    txt_file.write_text(txt_content)
    sample_files['txt'] = str(txt_file)
    print("✅ Created sample.txt")
    
    # 2. HTML file
    html_file = temp_dir / "sample.html"
    html_content = """<!DOCTYPE html>
<html>
<head>
    <title>Sample HTML Document</title>
</head>
<body>
    <h1>Sample HTML Document</h1>
    <p>This is a sample HTML document for testing docling.</p>
    <ul>
        <li>HTML formatting</li>
        <li>Lists and tables</li>
        <li>Links and images</li>
    </ul>
    <table border="1">
        <tr><th>Column 1</th><th>Column 2</th></tr>
        <tr><td>Data 1</td><td>Data 2</td></tr>
    </table>
</body>
</html>"""
    html_file.write_text(html_content)
    sample_files['html'] = str(html_file)
    print("✅ Created sample.html")
    
    # 3. Markdown file
    md_file = temp_dir / "sample.md"
    md_content = """# Sample Markdown Document

This is a sample markdown document for testing docling.

## Features

- **Bold text**
- *Italic text*
- `Code snippets`
- Lists and tables

### Code Example

```python
def hello_world():
    print("Hello, World!")
```

### Table

| Column 1 | Column 2 |
|----------|----------|
| Data 1   | Data 2   |
| Data 3   | Data 4   |
"""
    md_file.write_text(md_content)
    sample_files['md'] = str(md_file)
    print("✅ Created sample.md")
    
    # 4. Create a simple PDF (using reportlab if available)
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_file = temp_dir / "sample.pdf"
        c = canvas.Canvas(str(pdf_file), pagesize=letter)
        c.drawString(100, 750, "Sample PDF Document")
        c.drawString(100, 700, "This is a sample PDF document for testing docling.")
        c.drawString(100, 650, "Features:")
        c.drawString(120, 620, "- PDF text extraction")
        c.drawString(120, 590, "- Table recognition")
        c.drawString(120, 560, "- Image processing")
        c.save()
        sample_files['pdf'] = str(pdf_file)
        print("✅ Created sample.pdf")
    except ImportError:
        print("⚠️  reportlab not available, skipping PDF creation")
    
    # 5. Create DOCX file (using python-docx)
    try:
        from docx import Document
        
        docx_file = temp_dir / "sample.docx"
        doc = Document()
        doc.add_heading('Sample DOCX Document', 0)
        doc.add_paragraph('This is a sample DOCX document for testing docling.')
        
        doc.add_heading('Features', level=1)
        doc.add_paragraph('Docling supports various Word document features:')
        
        # Add a list
        doc.add_paragraph('• Text formatting', style='List Bullet')
        doc.add_paragraph('• Tables and images', style='List Bullet')
        doc.add_paragraph('• Headers and footers', style='List Bullet')
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Column 1'
        table.cell(0, 1).text = 'Column 2'
        table.cell(1, 0).text = 'Data 1'
        table.cell(1, 1).text = 'Data 2'
        
        doc.save(str(docx_file))
        sample_files['docx'] = str(docx_file)
        print("✅ Created sample.docx")
    except ImportError:
        print("⚠️  python-docx not available, skipping DOCX creation")
    
    # 6. Create XLSX file (using openpyxl)
    try:
        from openpyxl import Workbook
        
        xlsx_file = temp_dir / "sample.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = "Sample Sheet"
        
        # Add headers
        ws['A1'] = 'Name'
        ws['B1'] = 'Age'
        ws['C1'] = 'City'
        
        # Add data
        ws['A2'] = 'John Doe'
        ws['B2'] = 30
        ws['C2'] = 'New York'
        
        ws['A3'] = 'Jane Smith'
        ws['B3'] = 25
        ws['C3'] = 'Los Angeles'
        
        wb.save(str(xlsx_file))
        sample_files['xlsx'] = str(xlsx_file)
        print("✅ Created sample.xlsx")
    except ImportError:
        print("⚠️  openpyxl not available, skipping XLSX creation")
    
    return sample_files, temp_dir

def test_document_conversion(sample_files):
    """Test document conversion with various formats"""
    print("\n" + "=" * 60)
    print("TESTING DOCUMENT CONVERSION")
    print("=" * 60)
    
    converter = DocumentConverter()
    results = {}
    
    for format_type, file_path in sample_files.items():
        try:
            print(f"\n🔄 Testing {format_type.upper()} format: {file_path}")
            
            # Convert document
            result = converter.convert(file_path)
            
            # Export to markdown
            markdown_output = result.document.export_to_markdown()
            
            # Export to text
            text_output = result.document.export_to_text()
            
            results[format_type] = {
                'success': True,
                'markdown_length': len(markdown_output),
                'text_length': len(text_output),
                'markdown_preview': markdown_output[:200] + "..." if len(markdown_output) > 200 else markdown_output
            }
            
            print(f"✅ {format_type.upper()} conversion: SUCCESS")
            print(f"   Markdown length: {len(markdown_output)} chars")
            print(f"   Text length: {len(text_output)} chars")
            print(f"   Preview: {markdown_output[:100]}...")
            
        except Exception as e:
            print(f"❌ {format_type.upper()} conversion failed: {e}")
            results[format_type] = {'success': False, 'error': str(e)}
    
    return results

def test_cli_interface(sample_files):
    """Test docling CLI interface"""
    print("\n" + "=" * 60)
    print("TESTING CLI INTERFACE")
    print("=" * 60)
    
    cli_results = {}
    
    for format_type, file_path in sample_files.items():
        try:
            print(f"\n🔄 Testing CLI with {format_type.upper()}: {file_path}")
            
            # Test basic CLI command
            result = subprocess.run([sys.executable, "-m", "docling", file_path], 
                                  capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0:
                cli_results[format_type] = {
                    'success': True,
                    'output_length': len(result.stdout),
                    'preview': result.stdout[:200] + "..." if len(result.stdout) > 200 else result.stdout
                }
                print(f"✅ CLI {format_type.upper()}: SUCCESS")
                print(f"   Output length: {len(result.stdout)} chars")
            else:
                cli_results[format_type] = {
                    'success': False,
                    'error': result.stderr
                }
                print(f"❌ CLI {format_type.upper()} failed: {result.stderr}")
                
        except subprocess.TimeoutExpired:
            print(f"⏰ CLI {format_type.upper()} timed out")
            cli_results[format_type] = {'success': False, 'error': 'Timeout'}
        except Exception as e:
            print(f"❌ CLI {format_type.upper()} error: {e}")
            cli_results[format_type] = {'success': False, 'error': str(e)}
    
    return cli_results

def test_supported_formats():
    """Test docling's supported input formats"""
    print("\n" + "=" * 60)
    print("TESTING SUPPORTED FORMATS")
    print("=" * 60)
    
    try:
        from docling.datamodel.base_models import InputFormat
        
        print("📋 Supported input formats:")
        for format_type in InputFormat:
            print(f"   ✅ {format_type.value}")
        
        return True
        
    except Exception as e:
        print(f"❌ Failed to get supported formats: {e}")
        return False

def cleanup(temp_dir):
    """Clean up temporary files"""
    print("\n" + "=" * 60)
    print("CLEANUP")
    print("=" * 60)
    
    try:
        import shutil
        shutil.rmtree(temp_dir)
        print(f"✅ Cleaned up temp directory: {temp_dir}")
    except Exception as e:
        print(f"⚠️  Cleanup warning: {e}")

def main():
    """Main test function"""
    print("🚀 DOCLING COMPREHENSIVE TEST SUITE")
    print("Testing pip install and functionality across various formats")
    
    # Test 1: Pip installation
    install_success = test_pip_install()
    
    # Test 2: Import testing
    import_success = test_import()
    
    # Test 3: Supported formats
    formats_success = test_supported_formats()
    
    if not (install_success and import_success):
        print("\n❌ Basic installation/import failed. Stopping tests.")
        return
    
    # Test 4: Create sample files
    sample_files, temp_dir = create_sample_files()
    
    # Test 5: Document conversion
    conversion_results = test_document_conversion(sample_files)
    
    # Test 6: CLI interface
    cli_results = test_cli_interface(sample_files)
    
    # Test 7: Cleanup
    cleanup(temp_dir)
    
    # Summary
    print("\n" + "=" * 60)
    print("TEST SUMMARY")
    print("=" * 60)
    
    print(f"✅ Pip installation: {'PASS' if install_success else 'FAIL'}")
    print(f"✅ Import testing: {'PASS' if import_success else 'FAIL'}")
    print(f"✅ Supported formats: {'PASS' if formats_success else 'FAIL'}")
    
    print("\n📊 Document Conversion Results:")
    for format_type, result in conversion_results.items():
        status = "PASS" if result['success'] else "FAIL"
        print(f"   {format_type.upper()}: {status}")
    
    print("\n📊 CLI Interface Results:")
    for format_type, result in cli_results.items():
        status = "PASS" if result['success'] else "FAIL"
        print(f"   {format_type.upper()}: {status}")
    
    print("\n🎉 Test suite completed!")

if __name__ == "__main__":
    main()
